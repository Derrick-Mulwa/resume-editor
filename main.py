import docx
import pendulum
from docx.shared import Pt


resume = r"cover letter.docx"
doc = docx.Document(resume)


def get_details():
    employer = input("Enter employer's name/title     :")
    company  = input("Enter company's name            :")
    address  = input("Enter company's address         :")
    title    = input("Enter letter's title (RE: )     :")
    position = input("Enter application position/role :")

    wrong_input = input(f"Confirm your details.:\n1. Employer: {employer}\n"
          f"2. Company : {company}\n"
          f"3. Address : {address}\n"
          f"4: Title   : {title}\n"
          f"5. Position: {position}\n"
          f"Enter the number with a wrong value or any other key to continue: ")

    while wrong_input in ["1", "2", "3", "4", "5"]:
        if wrong_input == "1":
            employer = input("Enter employer's name/title     :")

        elif wrong_input == "2":
            company = input("Enter company's name            :")

        elif wrong_input == "3":
            address = input("Enter company's address         :")

        elif wrong_input == "4":
            title = input("Enter letter's title (RE: )     :")

        elif wrong_input == "5":
            position = input("Enter application position/role :")

        wrong_input = input(f"\n\nConfirm your details:\n1. Employer: {employer}\n"
                            f"2. Company : {company}\n"
                            f"3. Address : {address}\n"
                            f"4: Title   : {title}\n"
                            f"5. Position: {position}\n"
                            f"Enter the number with a wrong value or any other key to continue: ")

    return employer, company, address, title, position


def edit_date():
    edited_date = f"{pendulum.now().to_formatted_date_string()}."
    para = doc.paragraphs[4].add_run(edited_date)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)


def edit_employer(employers_name):
    employers_name_6 = f"To The {employers_name},"
    employers_name_10 = f"Dear {employers_name.capitalize()},"

    para = doc.paragraphs[6].add_run(employers_name_6)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)

    para = doc.paragraphs[10].add_run(employers_name_10)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)


def edit_company(company):
    company = f"{company},"
    para = doc.paragraphs[7].add_run(company)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)


def edit_address(address):
    address = f"{address}."
    para = doc.paragraphs[8].add_run(address)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)


def edit_title(title):
    title = "RE: "+title.upper()
    para = doc.paragraphs[12].add_run(title)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)
    para.underline = True
    para.bold = True


def edit_position_and_company(position, company):
    v = f"I am writing to express my strong interest in the {position.title()} position at {company}. " \
        f"With a combination of hands-on experience and a deep understanding of data analysis techniques and tools, " \
        f"I am confident that I have the skills and knowledge to thrive in this role. " \
        f"I have a track record of successfully analyzing and interpreting data to provide valuable insights " \
        f"and recommendations to clients and teams. I am confident that my skills and experience make me a strong " \
        f"candidate for this role at {company}."
    para = doc.paragraphs[14].add_run(v)
    para.font.name = 'Roboto'
    para.font.size = Pt(12)


if __name__ == "__main__":
    exit = False
    while exit is False:
        action = input("Welcome to Firefly resume editor. \n"
              "Press 1 to edit a new resume and any other key to exit: ")
        if action == "1":
            data = get_details()
            employer, company, address, title, position = data[0], data[1],data[2],data[3],data[4]

            edit_date()
            edit_employer(employer)
            edit_company(company)
            edit_address(address)
            edit_title(title)
            edit_position_and_company(position, company)

            doc.save(f'Resumes\Derrick Mulwa\'s Application Letter for {company}.docx')

            print(f'Resume saved as Derrick Mulwa\'s Application Letter for {company}\n\n')

        else:
            x = input("Thank you for using our service. ")
            exit = True
