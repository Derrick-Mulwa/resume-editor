import shutil
import docx
import os
import pendulum

resume = r"cover letter.docx"
doc = docx.Document(resume)
docxParagraphs = []

for para in doc.paragraphs:
    docxParagraphs.append(para.text)


date = docxParagraphs[4]
employer = docxParagraphs[6]
company = docxParagraphs[7]
address = docxParagraphs[8]
dear_employer = docxParagraphs[10]
title = docxParagraphs[12]
position_and_company = docxParagraphs[14]

print(employer)
def edit_date():
    edited_date = f"{pendulum.now().to_formatted_date_string()},"
    doc.paragraphs[4].text = edited_date


def edit_employer(employers_name):
    doc.paragraphs[6].text = f" To {employers_name},"

def edit_company(company):
    doc.paragraphs[6].text = f" To {employers_name},"




# doc.save('CV_edited.docx')
